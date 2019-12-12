from konlpy.tag import Okt
from collections import Counter
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
import pandas as pd
import sys
import copy
import re



def get_tags(text, ntags=10):
    spliter = Okt()
    # konlpy의 Twitter객체
    nouns = spliter.nouns(text)
    # nouns 함수를 통해서 text에서 명사만 분리/추출
    count = Counter(nouns)
    # Counter객체를 생성하고 참조변수 nouns할당
    return_list = []  # 명사 빈도수 저장할 변수
    for n, c in count.most_common(ntags):
        temp = {'tag': n, 'count': c}
        if len(temp['tag'])>=2 :
            return_list.append(temp)
    # most_common 메소드는 정수를 입력받아 객체 안의 명사중 빈도수
    # 큰 명사부터 순서대로 입력받은 정수 갯수만큼 저장되어있는 객체 반환
    # 명사와 사용된 갯수를 return_list에 저장합니다.
    return return_list


def main():
    text_file_name ='/Users/eunsung/PycharmProjects/ProjectFile/add.txt'
    # 분석할 파일
    noun_count = 10
    # 최대 많은 빈도수 부터 10개 명사 추출
    output_file_name = '/Users/eunsung/PycharmProjects/ProjectFile/counter.txt'
    # counter.txt 에 저장
    open_text_file = open(text_file_name, 'r', -1, "utf-8")
    # 분석할 파일을 open
    text = open_text_file.read()  # 파일을 읽습니다.
    tags = get_tags(text, noun_count)  # get_tags 함수 실행
    open_text_file.close()  # 파일 close
    open_output_file = open(output_file_name, 'w', -1, "utf-8")
    # 결과로 쓰일 count.txt 열기
    for tag in tags:
        noun = tag['tag']
        count = tag['count']
        open_output_file.write('{} {}\n'.format(noun, count))
    # 결과 저장
    open_output_file.close()




def highlight_text(filename, find):
    doc = Document(filename)
    for paragraph in doc.paragraphs:
        if find in paragraph.text:
            for run in paragraph.runs:
                if find in run.text:
                    x = run.text.split(find)
                    run.clear()
                    for i in range(len(x) - 1):
                        run.add_text(x[i])
                        run.add_text(find)
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    doc.save('/Users/eunsung/PycharmProjects/ProjectFile/blog8.docx')
    return 1

def highlight2(filename, find):

    source = filename
    phrase = find

    doc = Document(source)

    for para in doc.paragraphs :
        start = para.text.find(phrase)
        if start > -1 :
            pre = para.text[:start]
            post = para.text[start+len(phrase):]
            para.text = pre
            para.add_run(phrase)
            para.runs[1].font.highlight_color = WD_COLOR_INDEX.YELLOW
            para.add_run(post)

    doc.save('/Users/eunsung/PycharmProjects/ProjectFile/blog7.docx')

def Highlighting(infiledir, keyword):
    doc = Document(infiledir)
    #docx 파일을 불러옴.
    p1_text = doc.paragraphs[0].text
    doc.paragraphs[0].clear()
    p2 = doc.add_paragraph()
    substrings = p1_text.split(keyword)
    #keyword 위치를 찾아 highlighting시켜서 다시씀.
    for substring in substrings[:-1]:
        p2.add_run(substring)
        font = p2.add_run(keyword).font
        font.highlight_color = WD_COLOR_INDEX.YELLOW
    p2.add_run(substrings[-1])
    doc.save(infiledir)


def newHilighting(docxFileName,list):
    doc = Document(docxFileName)


    for paragraph in doc.paragraphs:
        for target in list:
            if target in paragraph.text:  # it is worth checking in detail ...

                currRuns = copy.copy(paragraph.runs)  # deep copy as we delete/clear the object
                paragraph.clear()

                for run in currRuns:
                    if target in run.text:
                        words = re.split('(\W)', run.text)  # split into words in order to be able to color only one
                        for word in words:
                            if word == target:
                                newRun = paragraph.add_run(word)
                                newRun.font.highlight_color = WD_COLOR_INDEX.PINK
                            else:
                                newRun = paragraph.add_run(word)
                                newRun.font.highlight_color = None
                    else:  # our target is not in it so we add it unchanged
                        paragraph.runs.append(run)

    doc.save('output.docx')


if __name__ == '__main__':
    main()




